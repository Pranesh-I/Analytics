from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
import pdfplumber
from docx import Document


def get_extension(file_path: str) -> str:
    return Path(file_path).suffix.lower().replace(".", "")


def _excel_engine(ext: str) -> str:
    if ext == "xls":
        return "xlrd"
    return "openpyxl"


def read_excel_file(file_path: str) -> Dict[str, Any]:
    """
    Read Excel/CSV files as raw sheets.
    IMPORTANT: header=None so we can detect the real header rows ourselves.
    """
    ext = get_extension(file_path)

    if ext in {"xls", "xlsx"}:
        engine = _excel_engine(ext)
        xls = pd.ExcelFile(file_path, engine=engine)
        sheets = {}

        for sheet_name in xls.sheet_names:
            df = pd.read_excel(
                file_path,
                sheet_name=sheet_name,
                header=None,
                engine=engine,
            )
            sheets[sheet_name] = df

        return {"type": "excel", "sheets": sheets}

    if ext == "csv":
        df = pd.read_csv(file_path, header=None)
        return {"type": "csv", "sheets": {"Sheet1": df}}

    raise ValueError(f"Unsupported tabular file type: {ext}")


def read_pdf_file(file_path: str) -> Dict[str, Any]:
    pages_data: List[Dict[str, Any]] = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            pages_data.append(
                {
                    "page": i,
                    "text": page.extract_text(),
                    "tables": page.extract_tables(),
                }
            )

    return {"type": "pdf", "pages": pages_data}


def read_word_file(file_path: str) -> Dict[str, Any]:
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        tables.append(table_data)

    return {"type": "docx", "paragraphs": paragraphs, "tables": tables}


def read_uploaded_file(file_path: str) -> Dict[str, Any]:
    ext = get_extension(file_path)

    if ext in {"xls", "xlsx", "csv"}:
        return read_excel_file(file_path)

    if ext == "pdf":
        return read_pdf_file(file_path)

    if ext in {"doc", "docx"}:
        return read_word_file(file_path)

    raise ValueError(f"Unsupported file type: {ext}")
